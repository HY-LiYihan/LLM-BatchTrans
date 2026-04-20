from __future__ import annotations

import json
import logging
from collections import Counter, defaultdict
from dataclasses import asdict
from datetime import datetime
from pathlib import Path

import pandas as pd
from docx import Document
from docx.shared import Pt

from .config import PipelineConfig
from .models import ArtifactManifest, Block, ExtractionReport, TranslationResult
from .review import build_review_markdown
from .text_utils import normalize_term


class ArtifactExporter:
    def __init__(self, logger: logging.Logger) -> None:
        self.logger = logger

    def export(
        self,
        *,
        extraction_report: ExtractionReport,
        translations: list[TranslationResult],
        output_dir: Path,
        pdf_path: Path,
        config: PipelineConfig,
    ) -> ArtifactManifest:
        output_dir.mkdir(parents=True, exist_ok=True)

        extracted_csv = output_dir / f"{pdf_path.stem}_提取文本.csv"
        extracted_xlsx = output_dir / f"{pdf_path.stem}_提取文本.xlsx"
        bilingual_csv = output_dir / f"{pdf_path.stem}_英中对照.csv"
        bilingual_xlsx = output_dir / f"{pdf_path.stem}_英中对照.xlsx"
        glossary_xlsx = output_dir / f"{pdf_path.stem}_术语对照表.xlsx"
        docx_path = output_dir / f"{pdf_path.stem}_中文译文.docx"
        review_path = output_dir / f"{pdf_path.stem}_自检报告.md"
        raw_json_path = output_dir / f"{pdf_path.stem}_result.json"
        config_json_path = output_dir / "run_config.json"

        extracted_df = self._build_extraction_dataframe(extraction_report)
        bilingual_df = self._build_bilingual_dataframe(translations)
        glossary_df = self._build_glossary_dataframe(translations)

        extracted_df.to_csv(extracted_csv, index=False, encoding="utf-8-sig")
        extracted_df.to_excel(extracted_xlsx, index=False)
        bilingual_df.to_csv(bilingual_csv, index=False, encoding="utf-8-sig")
        bilingual_df.to_excel(bilingual_xlsx, index=False)
        glossary_df.to_excel(glossary_xlsx, index=False)
        self._write_docx(translations, docx_path, pdf_path.stem)
        review_path.write_text(build_review_markdown(translations), encoding="utf-8")
        raw_json_path.write_text(
            json.dumps([asdict(item) for item in translations], ensure_ascii=False, indent=2),
            encoding="utf-8",
        )
        config_json_path.write_text(
            json.dumps(
                {
                    "pdf_path": str(pdf_path),
                    "generated_at": datetime.now().isoformat(),
                    "config": config.to_safe_dict(),
                    "skipped_pages": extraction_report.skipped_pages,
                },
                ensure_ascii=False,
                indent=2,
            ),
            encoding="utf-8",
        )

        manifest = ArtifactManifest(
            extracted_csv=extracted_csv,
            extracted_xlsx=extracted_xlsx,
            bilingual_csv=bilingual_csv,
            bilingual_xlsx=bilingual_xlsx,
            glossary_xlsx=glossary_xlsx,
            docx=docx_path,
            review=review_path,
            raw_json=raw_json_path,
            config_json=config_json_path,
        )
        self.logger.info("Artifacts generated:")
        for name, path in asdict(manifest).items():
            self.logger.info("  %s -> %s", name, path)
        return manifest

    def _build_extraction_dataframe(self, report: ExtractionReport) -> pd.DataFrame:
        rows = [
            {
                "Page": block.page,
                "Block ID": block.block_id,
                "Source Text": block.source_text,
                "Word Count": block.source_word_count,
                "Character Count": block.source_char_count,
                "Extraction Backend": block.extraction_backend,
                "Skipped Pages": ", ".join(str(page) for page in report.skipped_pages),
            }
            for block in report.blocks
        ]
        return pd.DataFrame(rows)

    def _build_bilingual_dataframe(self, translations: list[TranslationResult]) -> pd.DataFrame:
        rows = [
            {
                "Page": item.page,
                "Block ID": item.block_id,
                "Source Text": item.source_text,
                "Chinese Translation": item.translation,
                "Status": item.status,
                "Cache Hit": item.cache_hit,
                "Attempts": item.attempts,
                "Elapsed Seconds": round(item.elapsed_seconds, 2),
                "Error Message": item.error_message,
            }
            for item in translations
        ]
        return pd.DataFrame(rows)

    def _build_glossary_dataframe(self, translations: list[TranslationResult]) -> pd.DataFrame:
        counter: Counter[tuple[str, str]] = Counter()
        pages: defaultdict[tuple[str, str], set[int]] = defaultdict(set)
        examples: dict[tuple[str, str], str] = {}

        for result in translations:
            for pair in result.term_pairs:
                source = normalize_term(pair.source)
                target = pair.target.strip()
                if not source or not target:
                    continue
                key = (source, target)
                counter[key] += 1
                pages[key].add(result.page)
                examples.setdefault(key, result.source_text[:180])

        rows = []
        for (source, target), frequency in counter.most_common():
            rows.append(
                {
                    "English Term": source,
                    "Chinese Term": target,
                    "Frequency": frequency,
                    "Pages": ", ".join(str(page) for page in sorted(pages[(source, target)])),
                    "Example": examples[(source, target)],
                }
            )
        return pd.DataFrame(rows)

    def _write_docx(self, translations: list[TranslationResult], output_path: Path, source_name: str) -> None:
        document = Document()
        normal_style = document.styles["Normal"]
        normal_style.font.name = "Songti SC"
        normal_style.font.size = Pt(11)

        document.add_heading(f"{source_name} 中文译文", level=1)
        current_page = None
        for item in translations:
            if item.page != current_page:
                current_page = item.page
                document.add_heading(f"第 {current_page} 页", level=2)
            paragraph = document.add_paragraph()
            if item.status == "ok":
                paragraph.add_run(item.translation)
            else:
                paragraph.add_run(f"[翻译失败] {item.block_id}\n{item.source_text}")
        document.save(str(output_path))
